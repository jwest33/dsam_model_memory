"""
Document parsing system for converting various file types into memories.
Supports PDF, DOCX, TXT, MD, HTML, CSV, JSON, and more.
"""

import os
import re
import json
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
import hashlib
import mimetypes

# Document parsing libraries
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import markdown
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of a document"""
    content: str
    metadata: Dict[str, Any]
    chunk_index: int
    total_chunks: int
    source_file: str
    file_type: str
    extraction_method: str
    word_count: int
    char_count: int
    
    def to_memory_text(self) -> str:
        """Convert chunk to memory-ready text"""
        # Add context about the chunk position
        position = f"Part {self.chunk_index + 1} of {self.total_chunks}"
        source = f"Source: {Path(self.source_file).name}"
        
        # Combine with content
        if self.total_chunks > 1:
            return f"[{position}] {self.content}\n[{source}]"
        else:
            return f"{self.content}\n[{source}]"


@dataclass
class ParsedDocument:
    """Represents a fully parsed document"""
    file_path: str
    file_type: str
    title: Optional[str]
    author: Optional[str]
    created_date: Optional[datetime]
    modified_date: Optional[datetime]
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    total_words: int
    total_chars: int
    extraction_errors: List[str]
    
    @property
    def success(self) -> bool:
        """Check if parsing was successful"""
        return len(self.chunks) > 0
    
    def get_summary(self) -> Dict[str, Any]:
        """Get parsing summary"""
        return {
            "file": self.file_path,
            "type": self.file_type,
            "title": self.title,
            "chunks": len(self.chunks),
            "words": self.total_words,
            "chars": self.total_chars,
            "errors": len(self.extraction_errors),
            "success": self.success
        }


class ChunkingStrategy:
    """Base class for document chunking strategies"""
    
    def __init__(self, 
                 max_chunk_size: int = 2000,
                 chunk_overlap: int = 200,
                 min_chunk_size: int = 100):
        self.max_chunk_size = max_chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
    
    def chunk(self, text: str) -> List[str]:
        """Split text into chunks"""
        raise NotImplementedError


class SentenceChunker(ChunkingStrategy):
    """Chunks text by sentences, respecting max size"""
    
    def chunk(self, text: str) -> List[str]:
        """Split text into sentence-based chunks"""
        # Split by sentence endings
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = []
        current_size = 0
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            # If single sentence exceeds max size, split it
            if sentence_size > self.max_chunk_size:
                # Save current chunk if any
                if current_chunk:
                    chunks.append(' '.join(current_chunk))
                    current_chunk = []
                    current_size = 0
                
                # Split long sentence by words
                words = sentence.split()
                word_chunk = []
                word_size = 0
                
                for word in words:
                    if word_size + len(word) + 1 > self.max_chunk_size:
                        if word_chunk:
                            chunks.append(' '.join(word_chunk))
                        word_chunk = [word]
                        word_size = len(word)
                    else:
                        word_chunk.append(word)
                        word_size += len(word) + 1
                
                if word_chunk:
                    chunks.append(' '.join(word_chunk))
            
            # Add sentence to current chunk if it fits
            elif current_size + sentence_size + 1 <= self.max_chunk_size:
                current_chunk.append(sentence)
                current_size += sentence_size + 1
            
            # Start new chunk
            else:
                if current_chunk:
                    chunks.append(' '.join(current_chunk))
                current_chunk = [sentence]
                current_size = sentence_size
        
        # Add remaining chunk
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        # Apply overlap for context continuity
        if self.chunk_overlap > 0 and len(chunks) > 1:
            overlapped_chunks = []
            for i, chunk in enumerate(chunks):
                if i > 0:
                    # Get overlap from previous chunk
                    prev_words = chunks[i-1].split()[-self.chunk_overlap:]
                    overlap_text = ' '.join(prev_words)
                    chunk = f"{overlap_text} {chunk}"
                overlapped_chunks.append(chunk)
            chunks = overlapped_chunks
        
        # Filter out chunks that are too small
        chunks = [c for c in chunks if len(c) >= self.min_chunk_size]
        
        return chunks


class ParagraphChunker(ChunkingStrategy):
    """Chunks text by paragraphs"""
    
    def chunk(self, text: str) -> List[str]:
        """Split text into paragraph-based chunks"""
        # Split by double newlines or common paragraph markers
        paragraphs = re.split(r'\n\s*\n', text)
        
        chunks = []
        current_chunk = []
        current_size = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            para_size = len(para)
            
            # If single paragraph exceeds max size, use sentence chunker
            if para_size > self.max_chunk_size:
                if current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = []
                    current_size = 0
                
                # Use sentence chunker for long paragraph
                sentence_chunker = SentenceChunker(
                    self.max_chunk_size,
                    self.chunk_overlap,
                    self.min_chunk_size
                )
                para_chunks = sentence_chunker.chunk(para)
                chunks.extend(para_chunks)
            
            # Add paragraph to current chunk if it fits
            elif current_size + para_size + 2 <= self.max_chunk_size:
                current_chunk.append(para)
                current_size += para_size + 2
            
            # Start new chunk
            else:
                if current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                current_chunk = [para]
                current_size = para_size
        
        # Add remaining chunk
        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))
        
        return [c for c in chunks if len(c) >= self.min_chunk_size]


class SemanticChunker(ChunkingStrategy):
    """Chunks text by semantic sections (headers, topics)"""
    
    def chunk(self, text: str) -> List[str]:
        """Split text into semantic chunks based on structure"""
        chunks = []
        
        # Look for markdown-style headers
        sections = re.split(r'^#{1,6}\s+', text, flags=re.MULTILINE)
        
        if len(sections) > 1:
            # Document has headers, use them
            for section in sections:
                if not section.strip():
                    continue
                
                # Check if section needs further chunking
                if len(section) > self.max_chunk_size:
                    para_chunker = ParagraphChunker(
                        self.max_chunk_size,
                        self.chunk_overlap,
                        self.min_chunk_size
                    )
                    section_chunks = para_chunker.chunk(section)
                    chunks.extend(section_chunks)
                else:
                    chunks.append(section.strip())
        else:
            # No headers found, fall back to paragraph chunking
            para_chunker = ParagraphChunker(
                self.max_chunk_size,
                self.chunk_overlap,
                self.min_chunk_size
            )
            chunks = para_chunker.chunk(text)
        
        return [c for c in chunks if len(c) >= self.min_chunk_size]


class DocumentParser:
    """Main document parsing class"""
    
    def __init__(self,
                 chunking_strategy: Optional[ChunkingStrategy] = None,
                 max_file_size: int = 50 * 1024 * 1024,  # 50MB
                 supported_types: Optional[List[str]] = None):
        """
        Initialize document parser
        
        Args:
            chunking_strategy: Strategy for splitting documents
            max_file_size: Maximum file size to process
            supported_types: List of supported MIME types
        """
        self.chunking_strategy = chunking_strategy or SemanticChunker()
        self.max_file_size = max_file_size
        
        # Default supported types
        self.supported_types = supported_types or [
            'text/plain',
            'text/markdown',
            'text/html',
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/json',
            'text/csv',
            'application/xml',
            'text/xml',
            'image/jpeg',
            'image/png',
            'image/gif',
            'image/bmp'
        ]
        
        # Extension to parser mapping
        self.parsers = {
            '.txt': self._parse_text,
            '.md': self._parse_markdown,
            '.markdown': self._parse_markdown,
            '.html': self._parse_html,
            '.htm': self._parse_html,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.json': self._parse_json,
            '.csv': self._parse_csv,
            '.xml': self._parse_xml,
            '.jpg': self._parse_image,
            '.jpeg': self._parse_image,
            '.png': self._parse_image,
            '.gif': self._parse_image,
            '.bmp': self._parse_image,
            '.py': self._parse_code,
            '.js': self._parse_code,
            '.java': self._parse_code,
            '.cpp': self._parse_code,
            '.c': self._parse_code,
            '.cs': self._parse_code,
            '.php': self._parse_code,
            '.rb': self._parse_code,
            '.go': self._parse_code,
            '.rs': self._parse_code,
            '.swift': self._parse_code,
            '.kt': self._parse_code,
            '.scala': self._parse_code,
            '.r': self._parse_code,
            '.m': self._parse_code,
            '.h': self._parse_code,
            '.hpp': self._parse_code,
            '.sh': self._parse_code,
            '.bash': self._parse_code,
            '.ps1': self._parse_code,
            '.sql': self._parse_code,
            '.yaml': self._parse_yaml,
            '.yml': self._parse_yaml,
            '.toml': self._parse_toml,
            '.ini': self._parse_ini,
            '.cfg': self._parse_ini,
            '.conf': self._parse_ini,
            '.log': self._parse_log,
        }
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a document file into chunks
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            ParsedDocument with chunks and metadata
        """
        file_path = Path(file_path)
        errors = []
        
        # Check file exists
        if not file_path.exists():
            return ParsedDocument(
                file_path=str(file_path),
                file_type="unknown",
                title=None,
                author=None,
                created_date=None,
                modified_date=None,
                chunks=[],
                metadata={},
                total_words=0,
                total_chars=0,
                extraction_errors=[f"File not found: {file_path}"]
            )
        
        # Check file size
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            return ParsedDocument(
                file_path=str(file_path),
                file_type="unknown",
                title=None,
                author=None,
                created_date=None,
                modified_date=None,
                chunks=[],
                metadata={"file_size": file_size},
                total_words=0,
                total_chars=0,
                extraction_errors=[f"File too large: {file_size} bytes (max: {self.max_file_size})"]
            )
        
        # Get file metadata
        file_stats = file_path.stat()
        created_date = datetime.fromtimestamp(file_stats.st_ctime)
        modified_date = datetime.fromtimestamp(file_stats.st_mtime)
        
        # Determine file type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        file_ext = file_path.suffix.lower()
        
        # Get appropriate parser
        parser_func = self.parsers.get(file_ext, self._parse_text)
        
        try:
            # Parse the document
            content, metadata = parser_func(file_path)
            
            # Extract title if not in metadata
            if 'title' not in metadata:
                metadata['title'] = file_path.stem
            
            # Chunk the content
            if content:
                chunks_text = self.chunking_strategy.chunk(content)
            else:
                chunks_text = []
                errors.append("No content extracted from file")
            
            # Create chunk objects
            chunks = []
            total_words = 0
            total_chars = 0
            
            for i, chunk_text in enumerate(chunks_text):
                word_count = len(chunk_text.split())
                char_count = len(chunk_text)
                total_words += word_count
                total_chars += char_count
                
                chunk = DocumentChunk(
                    content=chunk_text,
                    metadata={
                        "file_name": file_path.name,
                        "file_type": mime_type or file_ext,
                        "chunk_method": self.chunking_strategy.__class__.__name__,
                        **metadata
                    },
                    chunk_index=i,
                    total_chunks=len(chunks_text),
                    source_file=str(file_path),
                    file_type=file_ext,
                    extraction_method=parser_func.__name__,
                    word_count=word_count,
                    char_count=char_count
                )
                chunks.append(chunk)
            
            return ParsedDocument(
                file_path=str(file_path),
                file_type=mime_type or file_ext,
                title=metadata.get('title'),
                author=metadata.get('author'),
                created_date=created_date,
                modified_date=modified_date,
                chunks=chunks,
                metadata=metadata,
                total_words=total_words,
                total_chars=total_chars,
                extraction_errors=errors
            )
            
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            errors.append(str(e))
            
            return ParsedDocument(
                file_path=str(file_path),
                file_type=mime_type or file_ext,
                title=file_path.stem,
                author=None,
                created_date=created_date,
                modified_date=modified_date,
                chunks=[],
                metadata={"error": str(e)},
                total_words=0,
                total_chars=0,
                extraction_errors=errors
            )
    
    def _parse_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parse plain text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        return content, {}
    
    def _parse_markdown(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parse markdown file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        metadata = {}
        
        # Extract front matter if present
        if content.startswith('---'):
            parts = content.split('---', 2)
            if len(parts) >= 3:
                try:
                    import yaml
                    metadata = yaml.safe_load(parts[1])
                    content = parts[2]
                except:
                    pass
        
        # Convert markdown to plain text if library available
        if HAS_MARKDOWN:
            html = markdown.markdown(content)
            if HAS_BS4:
                soup = BeautifulSoup(html, 'html.parser')
                content = soup.get_text()
        
        return content, metadata
    
    def _parse_html(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parse HTML file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html = f.read()
        
        metadata = {}
        
        if HAS_BS4:
            soup = BeautifulSoup(html, 'html.parser')
            
            # Extract metadata
            title_tag = soup.find('title')
            if title_tag:
                metadata['title'] = title_tag.get_text()
            
            author_tag = soup.find('meta', {'name': 'author'})
            if author_tag:
                metadata['author'] = author_tag.get('content')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            content = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = '\n'.join(chunk for chunk in chunks if chunk)
        else:
            # Basic HTML stripping
            content = re.sub('<[^<]+?>', '', html)
        
        return content, metadata
    
    def _parse_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parse PDF file"""
        if not HAS_PDF:
            raise ImportError("PyPDF2 required for PDF parsing. Install with: pip install PyPDF2")
        
        content = []
        metadata = {}
        
        with open(file_path, 'rb') as f:
            pdf = PyPDF2.PdfReader(f)
            
            # Extract metadata
            if pdf.metadata:
                metadata['title'] = pdf.metadata.get('/Title', '')
                metadata['author'] = pdf.metadata.get('/Author', '')
                metadata['subject'] = pdf.metadata.get('/Subject', '')
            
            # Extract text from each page
            for page_num, page in enumerate(pdf.pages):
                try:
                    text = page.extract_text()
                    if text:
                        content.append(text)
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num}: {e}")
        
        return '\n'.join(content), metadata
    
    def _parse_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parse DOCX file"""
        if not HAS_DOCX:
            raise ImportError("python-docx required for DOCX parsing. Install with: pip install python-docx")
        
        doc = DocxDocument(file_path)
        
        # Extract metadata
        metadata = {}
        if doc.core_properties:
            metadata['title'] = doc.core_properties.title or ''
            metadata['author'] = doc.core_properties.author or ''
            metadata['subject'] = doc.core_properties.subject or ''
        
        # Extract text
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text:
                        row_text.append(cell.text)
                if row_text:
                    content.append(' | '.join(row_text))
        
        return '\n'.join(content), metadata
    
    def _parse_json(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parse JSON file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Convert JSON to readable text
        content = json.dumps(data, indent=2, ensure_ascii=False)
        
        metadata = {
            'type': 'json',
            'keys': list(data.keys()) if isinstance(data, dict) else None,
            'length': len(data) if isinstance(data, (list, dict)) else None
        }
        
        return content, metadata
    
    def _parse_csv(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parse CSV file"""
        if HAS_PANDAS:
            df = pd.read_csv(file_path)
            
            # Convert to readable text
            content = df.to_string()
            
            metadata = {
                'type': 'csv',
                'columns': list(df.columns),
                'rows': len(df),
                'shape': df.shape
            }
        else:
            # Basic CSV parsing
            import csv
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f)
                rows = list(reader)
            
            if rows:
                metadata = {
                    'type': 'csv',
                    'columns': rows[0] if rows else [],
                    'rows': len(rows) - 1
                }
                content = '\n'.join([', '.join(row) for row in rows])
            else:
                content = ""
                metadata = {'type': 'csv', 'rows': 0}
        
        return content, metadata
    
    def _parse_xml(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parse XML file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            xml_content = f.read()
        
        if HAS_BS4:
            soup = BeautifulSoup(xml_content, 'xml')
            content = soup.get_text()
        else:
            # Basic XML stripping
            content = re.sub('<[^<]+?>', '', xml_content)
        
        metadata = {'type': 'xml'}
        return content, metadata
    
    def _parse_yaml(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parse YAML file"""
        try:
            import yaml
            with open(file_path, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)
            
            # Convert to readable text
            content = yaml.dump(data, default_flow_style=False)
            metadata = {
                'type': 'yaml',
                'keys': list(data.keys()) if isinstance(data, dict) else None
            }
        except ImportError:
            # Fall back to text parsing
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            metadata = {'type': 'yaml'}
        
        return content, metadata
    
    def _parse_toml(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parse TOML file"""
        try:
            import toml
            with open(file_path, 'r', encoding='utf-8') as f:
                data = toml.load(f)
            
            content = toml.dumps(data)
            metadata = {
                'type': 'toml',
                'sections': list(data.keys()) if isinstance(data, dict) else None
            }
        except ImportError:
            # Fall back to text parsing
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            metadata = {'type': 'toml'}
        
        return content, metadata
    
    def _parse_ini(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parse INI/config file"""
        import configparser
        
        config = configparser.ConfigParser()
        config.read(file_path)
        
        content = []
        sections = []
        
        for section in config.sections():
            sections.append(section)
            content.append(f"[{section}]")
            for key, value in config.items(section):
                content.append(f"{key} = {value}")
            content.append("")
        
        metadata = {
            'type': 'ini',
            'sections': sections
        }
        
        return '\n'.join(content), metadata
    
    def _parse_code(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parse source code file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Extract basic metadata
        lines = content.splitlines()
        metadata = {
            'type': 'code',
            'language': file_path.suffix[1:],
            'lines': len(lines),
            'size': len(content)
        }
        
        # Look for common patterns
        imports = []
        functions = []
        classes = []
        
        for line in lines:
            # Python/JS imports
            if re.match(r'^(import|from|const|let|var|require)', line):
                imports.append(line.strip())
            # Function definitions
            elif re.match(r'^(def |function |func |fn |sub )', line):
                functions.append(line.strip())
            # Class definitions
            elif re.match(r'^(class |interface |struct |enum )', line):
                classes.append(line.strip())
        
        if imports:
            metadata['imports'] = imports[:10]  # First 10 imports
        if functions:
            metadata['functions'] = functions[:10]  # First 10 functions
        if classes:
            metadata['classes'] = classes[:10]  # First 10 classes
        
        return content, metadata
    
    def _parse_log(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parse log file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
        
        # Analyze log patterns
        metadata = {
            'type': 'log',
            'lines': len(lines)
        }
        
        # Look for common log levels
        levels = {'ERROR': 0, 'WARNING': 0, 'INFO': 0, 'DEBUG': 0}
        for line in lines:
            for level in levels:
                if level in line.upper():
                    levels[level] += 1
        
        metadata['levels'] = levels
        
        # Join lines for content
        content = ''.join(lines)
        
        return content, metadata
    
    def _parse_image(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parse image file using OCR if available"""
        metadata = {
            'type': 'image',
            'format': file_path.suffix[1:].upper()
        }
        
        if HAS_OCR:
            try:
                # Open image
                image = Image.open(file_path)
                
                # Extract text using OCR
                content = pytesseract.image_to_string(image)
                
                # Add image metadata
                metadata['size'] = image.size
                metadata['mode'] = image.mode
                
                if not content.strip():
                    content = f"[Image: {file_path.name}]"
                    metadata['ocr_text'] = False
                else:
                    metadata['ocr_text'] = True
                    
            except Exception as e:
                logger.warning(f"OCR failed for {file_path}: {e}")
                content = f"[Image: {file_path.name}]"
                metadata['ocr_error'] = str(e)
        else:
            content = f"[Image: {file_path.name}]"
            metadata['ocr_available'] = False
        
        return content, metadata
    
    def parse_directory(self, dir_path: str, 
                       recursive: bool = True,
                       extensions: Optional[List[str]] = None) -> List[ParsedDocument]:
        """
        Parse all documents in a directory
        
        Args:
            dir_path: Path to directory
            recursive: Whether to process subdirectories
            extensions: List of file extensions to process (e.g., ['.txt', '.pdf'])
            
        Returns:
            List of ParsedDocument objects
        """
        dir_path = Path(dir_path)
        
        if not dir_path.is_dir():
            logger.error(f"Not a directory: {dir_path}")
            return []
        
        # Get all files
        if recursive:
            pattern = '**/*'
        else:
            pattern = '*'
        
        files = []
        for file_path in dir_path.glob(pattern):
            if file_path.is_file():
                # Check extension filter
                if extensions:
                    if file_path.suffix.lower() not in extensions:
                        continue
                
                # Check if supported
                if file_path.suffix.lower() in self.parsers:
                    files.append(file_path)
        
        # Parse all files
        results = []
        for file_path in files:
            logger.info(f"Parsing: {file_path}")
            doc = self.parse(str(file_path))
            results.append(doc)
        
        return results


# Convenience functions
def parse_document(file_path: str, 
                  chunking_strategy: Optional[ChunkingStrategy] = None) -> ParsedDocument:
    """
    Parse a single document
    
    Args:
        file_path: Path to document
        chunking_strategy: Optional chunking strategy
        
    Returns:
        ParsedDocument object
    """
    parser = DocumentParser(chunking_strategy=chunking_strategy)
    return parser.parse(file_path)


def parse_documents(file_paths: List[str],
                   chunking_strategy: Optional[ChunkingStrategy] = None) -> List[ParsedDocument]:
    """
    Parse multiple documents
    
    Args:
        file_paths: List of file paths
        chunking_strategy: Optional chunking strategy
        
    Returns:
        List of ParsedDocument objects
    """
    parser = DocumentParser(chunking_strategy=chunking_strategy)
    results = []
    
    for file_path in file_paths:
        doc = parser.parse(file_path)
        results.append(doc)
    
    return results